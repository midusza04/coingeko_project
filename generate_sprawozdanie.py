"""
Generuje spójny raport akademicki SPRAWOZDANIE.docx.
Źródło treści: SPRAWOZDANIE.md + DOCUMENTATION.md + .docs/

Uruchomienie:
    uv run python generate_sprawozdanie.py
"""

import base64
import time
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Inches

OUTPUT = "SPRAWOZDANIE.docx"
TMP_DIR = Path("_tmp_diagrams")


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int, color: str = "1F3864") -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 11, space_before: int = 0, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(11)


def add_code(doc: Document, code: str) -> None:
    for line in code.rstrip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F4F6")
        pPr.append(shd)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        pPr.append(ind)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("1E1B4B")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_table(doc: Document, headers: list, rows: list,
              col_widths: list | None = None,
              header_bg: str = "1F3864") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_data in enumerate(rows):
        bg = "EEF2FF" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Cm(col_widths[i])

    doc.add_paragraph()


def render_mermaid(code: str, name: str) -> Path | None:
    TMP_DIR.mkdir(exist_ok=True)
    encoded = base64.urlsafe_b64encode(code.encode()).decode()
    url = f"https://mermaid.ink/img/{encoded}?theme=default&bgColor=ffffff"
    print(f"  Renderuję diagram: {name}...")
    try:
        resp = requests.get(url, timeout=25)
        resp.raise_for_status()
        path = TMP_DIR / f"{name}.png"
        path.write_bytes(resp.content)
        print(f"  OK: {name} ({len(resp.content) // 1024} KB)")
        return path
    except Exception as exc:
        print(f"  Błąd {name}: {exc}")
        return None


def add_diagram(doc: Document, path: Path | None, caption: str, width_cm: float = 15.0) -> None:
    if path and path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        doc.add_paragraph()
    else:
        add_para(doc, f"[Diagram niedostępny: {caption}]", italic=True)


# ── diagram definitions ───────────────────────────────────────────────────────

DIAGRAMS = {
    "erd": """\
erDiagram
    cryptocurrencies {
        TEXT id PK
        TEXT symbol
        TEXT name
    }
    market_snapshots {
        INTEGER record_id PK
        TEXT crypto_id FK
        DATE snapshot_date
        REAL price_usd
        REAL market_cap
        REAL total_volume
    }
    market_current {
        INTEGER record_id PK
        TEXT crypto_id FK
        DATETIME collected_at
        REAL price_usd
        REAL market_cap
        REAL total_volume
        REAL high_24h
        REAL low_24h
        REAL price_change_percentage_24h
        REAL price_change_percentage_7d
        INTEGER market_cap_rank
        REAL ath
    }
    cryptocurrencies ||--o{ market_snapshots : ma_snapshoty
    cryptocurrencies ||--o{ market_current : ma_snapshoty_live
""",
    "architecture": """\
flowchart TD
    subgraph PRESENT["Warstwa prezentacji"]
        NB["Jupyter Notebook\\n11 etapow"]
        APP["Streamlit app.py\\n6 stron"]
    end
    subgraph STORAGE["Warstwa przechowywania"]
        DB[("SQLite\\ncrypto_market.db")]
    end
    subgraph COLLECT["Warstwa pozyskiwania"]
        F1["fetch_market_chart"]
        F2["fetch_markets_current"]
    end
    API["CoinGecko API v3"]
    API --> F1
    API --> F2
    F1 --> DB
    F2 --> DB
    DB --> NB
    DB --> APP
""",
    "sequence": """\
sequenceDiagram
    participant U as Uzytkownik
    participant S as Streamlit
    participant A as CoinGecko API
    participant D as SQLite
    U->>S: Fetch Historical Data
    loop 5 monet
        S->>A: GET /coins/id/market_chart
        A-->>S: JSON prices market_caps volumes
        S->>D: INSERT OR REPLACE market_snapshots
        S->>S: sleep 10s
    end
    S-->>U: 1826 wierszy zapisanych
""",
}


# ── build document ────────────────────────────────────────────────────────────

def build() -> None:
    print("Generuję diagramy...")
    rendered = {k: render_mermaid(v, k) for k, v in DIAGRAMS.items()}
    time.sleep(0.3)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ══ STRONA TYTUŁOWA ══════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    for text, size, bold, color in [
        ("POLITECHNIKA", 14, True, "1F3864"),
        ("Automatyka i Robotyka II Stopnia", 12, False, "404040"),
        ("Informatyka w Sterowaniu i Zarządzaniu", 12, False, "404040"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = RGBColor.from_string(color)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("System Analizy Rynku Kryptowalut")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("1F3864")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sprawozdanie z projektu")
    r.font.size = Pt(14)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Przedmiot: Zaawansowane Bazy Danych")
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Autorzy:\nMichał Dusza\nSzymon Bugajski\nMateusz Basiura")
    r.font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Maj 2026")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("555555")

    doc.add_page_break()

    # ══ SPIS TREŚCI ═════════════════════════════════════════════════════════
    add_heading(doc, "Spis treści", 1)
    toc = [
        "1. Wstęp i cel projektu",
        "2. Zakres funkcjonalny",
        "3. Zastosowane technologie",
        "4. Struktura repozytorium",
        "5. Projekt bazy danych",
        "6. Architektura systemu",
        "7. Implementacja — warstwa danych",
        "8. Implementacja — warstwa prezentacji",
        "9. Opis analiz i wizualizacji",
        "10. Dane zebrane w projekcie",
        "11. Instrukcja uruchomienia",
        "12. Napotkane problemy i rozwiązania",
        "13. Wnioski",
        "14. Literatura i źródła",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ══ 1. WSTĘP ════════════════════════════════════════════════════════════
    add_heading(doc, "1. Wstęp i cel projektu", 1)

    add_heading(doc, "1.1 Kontekst", 2)
    add_para(doc,
        "Rynek kryptowalut charakteryzuje się wyjątkowo wysoką zmiennością i generuje ogromną "
        "ilość danych w czasie rzeczywistym. Monitorowanie i analiza tych danych wymaga "
        "niezawodnej infrastruktury do ich przechowywania oraz narzędzi do interaktywnej "
        "eksploracji. Projekt odpowiada na to zapotrzebowanie, łącząc pobieranie danych "
        "przez REST API, relacyjne przechowywanie w bazie SQLite oraz interaktywne wizualizacje."
    )

    add_heading(doc, "1.2 Cel projektu", 2)
    add_para(doc, "Celem projektu było zaprojektowanie i zaimplementowanie systemu, który:")
    for g in [
        "Automatycznie pobiera historyczne i bieżące dane rynkowe kryptowalut z publicznego API CoinGecko.",
        "Przechowuje dane w znormalizowanej relacyjnej bazie danych SQLite z kluczami, ograniczeniami i indeksami.",
        "Prezentuje zebrane dane w interaktywnych wizualizacjach — w Jupyter i Streamlit.",
    ]:
        add_numbered(doc, g)

    add_heading(doc, "1.3 Motywacja wyboru tematu", 2)
    add_para(doc,
        "Kryptowaluty stanowią doskonały przypadek użycia dla systemów baz danych zorientowanych "
        "na szeregi czasowe: dane mają charakter cykliczny, wymagają idempotentnego zapisu, "
        "a analizy statystyczne (korelacje, zmienność, rozkłady) są naturalnymi operacjami."
    )

    doc.add_page_break()

    # ══ 2. ZAKRES ═════════════════════════════════════════════════════════════
    add_heading(doc, "2. Zakres funkcjonalny", 1)
    add_table(doc,
        ["Moduł", "Funkcjonalność"],
        [
            ["Pozyskiwanie danych", "365-dniowa historia (cena, kapitalizacja, wolumen) z CoinGecko API"],
            ["Pozyskiwanie danych", "Bieżący snapshot rynkowy (24h high/low, ATH, rank, supply)"],
            ["Baza danych", "3 tabele relacyjne: FK, UNIQUE, indeksy, normalizacja 3NF"],
            ["Szeregi czasowe", "Wykres liniowy z MA, skalą log, 6 filtrami"],
            ["Analiza ilościowa", "Bar / Box / Violin z agregacjami mean/max/min/std"],
            ["Dashboard rynkowy", "KPI, tabela, grouped bar, heatmapa, treemap"],
            ["Korelacja i zmienność", "Macierz korelacji, zmienność roczna, korelacja z BTC"],
        ],
        col_widths=[4.5, 12.5],
    )

    # ══ 3. TECHNOLOGIE ════════════════════════════════════════════════════════
    add_heading(doc, "3. Zastosowane technologie", 1)
    add_table(doc,
        ["Komponent", "Technologia", "Wersja", "Uzasadnienie"],
        [
            ["Język", "Python", "3.13", "Ekosystem data-science, SQLite w stdlib"],
            ["Menedżer pakietów", "uv", "0.10.9+", "Deterministyczny uv.lock"],
            ["Baza danych", "SQLite 3", "stdlib", "Zero-konfiguracyjna, pełny SQL + FK"],
            ["Źródło danych", "CoinGecko API v3", "public", "Darmowy dostęp, dane historyczne"],
            ["Notebook", "JupyterLab", "4.x", "ipywidgets + interaktywne analizy"],
            ["Wizualizacja", "Plotly", "6.7", "Interaktywne wykresy"],
            ["Aplikacja webowa", "Streamlit", "1.57", "Szybkie UI data-science"],
            ["Dane", "pandas", "2.2+", "pd.read_sql jako warstwa SQL↔Python"],
        ],
        col_widths=[4.0, 3.5, 2.5, 7.5],
    )

    doc.add_page_break()

    # ══ 4. STRUKTURA REPO ═════════════════════════════════════════════════════
    add_heading(doc, "4. Struktura repozytorium", 1)
    add_para(doc,
        "Repozytorium projektu zawiera kod źródłowy, bazę danych, dokumentację "
        "oraz wygenerowany raport. Poniżej opisano rolę każdego kluczowego pliku."
    )
    add_table(doc,
        ["Plik / folder", "Rola", "Status"],
        [
            ["app.py", "Główna aplikacja Streamlit — ETL + 6 stron UI", "główny"],
            ["crypto_market_analysis.ipynb", "Notebook analityczny — 11 etapów", "główny"],
            ["crypto_market.db", "Baza SQLite (tworzona automatycznie)", "główny"],
            ["SPRAWOZDANIE.md", "Raport akademicki (Markdown)", "dokumentacja"],
            ["DOCUMENTATION.md", "Dokumentacja techniczna (EN)", "dokumentacja"],
            [".docs/", "Dokumentacja techniczna (PL) — architektura, model, API", "dokumentacja"],
            ["main.py", "Prosty fetcher JSON (3 monety, bez bazy)", "legacy"],
            ["generate_sprawozdanie.py", "Generator tego raportu DOCX", "narzędzie"],
        ],
        col_widths=[5.5, 8.5, 3.5],
    )
    add_para(doc,
        "Przepływ danych w systemie:",
        bold=True, space_after=4,
    )
    add_code(doc,
        "CoinGecko API  -->  app.py / notebook  -->  crypto_market.db  -->  Plotly charts"
    )

    doc.add_page_break()

    # ══ 5. BAZA DANYCH ════════════════════════════════════════════════════════
    add_heading(doc, "5. Projekt bazy danych", 1)

    add_heading(doc, "5.1 Model konceptualny (ERD)", 2)
    add_para(doc,
        "Baza składa się z trzech tabel w modelu gwiazdy: wymiar cryptocurrencies "
        "oraz dwie tabele faktów — market_snapshots (historia) i market_current (live)."
    )
    add_diagram(doc, rendered.get("erd"), "Rys. 1. Diagram ERD — schemat relacyjnej bazy danych")

    add_heading(doc, "5.2 DDL — definicje tabel", 2)
    add_code(doc, """\
CREATE TABLE IF NOT EXISTS cryptocurrencies (
    id     TEXT PRIMARY KEY,
    symbol TEXT NOT NULL,
    name   TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS market_snapshots (
    record_id     INTEGER PRIMARY KEY AUTOINCREMENT,
    crypto_id     TEXT    NOT NULL,
    snapshot_date DATE    NOT NULL,
    price_usd     REAL,
    market_cap    REAL,
    total_volume  REAL,
    UNIQUE (crypto_id, snapshot_date),
    FOREIGN KEY (crypto_id) REFERENCES cryptocurrencies(id)
);

CREATE INDEX IF NOT EXISTS idx_snapshots_date
    ON market_snapshots(snapshot_date);

CREATE TABLE IF NOT EXISTS market_current (
    record_id                   INTEGER  PRIMARY KEY AUTOINCREMENT,
    crypto_id                   TEXT     NOT NULL,
    collected_at                DATETIME NOT NULL,
    price_usd                   REAL,  market_cap         REAL,
    total_volume                REAL,  high_24h           REAL,
    low_24h                     REAL,  price_change_24h   REAL,
    price_change_percentage_24h REAL,  price_change_percentage_7d REAL,
    market_cap_rank             INTEGER, circulating_supply REAL,
    total_supply                REAL,  max_supply         REAL,
    ath                         REAL,  ath_change_percentage REAL,
    FOREIGN KEY (crypto_id) REFERENCES cryptocurrencies(id)
);

CREATE INDEX IF NOT EXISTS idx_current_collected_at
    ON market_current(collected_at);""")

    add_heading(doc, "5.3 Opis tabel", 2)

    add_para(doc, "cryptocurrencies — słownik monet", bold=True, space_after=2)
    add_table(doc,
        ["Kolumna", "Typ", "Ograniczenia", "Opis"],
        [
            ["id", "TEXT", "PRIMARY KEY", "Identyfikator CoinGecko, np. bitcoin"],
            ["symbol", "TEXT", "NOT NULL", "Ticker, np. BTC"],
            ["name", "TEXT", "NOT NULL", "Pełna nazwa, np. Bitcoin"],
        ],
        col_widths=[3.0, 2.0, 3.5, 8.0],
    )

    add_para(doc, "market_snapshots — historia dzienna", bold=True, space_after=2)
    add_table(doc,
        ["Kolumna", "Typ", "Opis"],
        [
            ["record_id", "INTEGER PK", "Klucz surogatowy"],
            ["crypto_id", "TEXT FK", "→ cryptocurrencies.id"],
            ["snapshot_date", "DATE", "Data YYYY-MM-DD (UTC)"],
            ["price_usd", "REAL", "Cena zamknięcia USD"],
            ["market_cap", "REAL", "Kapitalizacja rynkowa USD"],
            ["total_volume", "REAL", "Wolumen 24h USD"],
        ],
        col_widths=[3.5, 2.5, 10.5],
    )

    add_para(doc, "market_current — bieżące snapshoty (append-only)", bold=True, space_after=2)
    add_para(doc,
        "Przechowuje bogate dane live z /coins/markets. Każde pobranie dodaje nowe wiersze "
        "z collected_at = UTC now — historia kolejnych pobrań."
    )

    add_heading(doc, "5.4 Ograniczenia integralności", 2)
    add_table(doc,
        ["Ograniczenie", "Tabela", "Definicja"],
        [
            ["PRIMARY KEY", "wszystkie", "id lub record_id"],
            ["FOREIGN KEY", "market_snapshots, market_current", "crypto_id → cryptocurrencies.id"],
            ["UNIQUE", "market_snapshots", "(crypto_id, snapshot_date)"],
            ["NOT NULL", "market_snapshots", "crypto_id, snapshot_date"],
            ["NOT NULL", "market_current", "crypto_id, collected_at"],
        ],
        col_widths=[3.5, 5.0, 8.0],
    )

    add_heading(doc, "5.5 Indeksy i normalizacja", 2)
    add_bullet(doc, "PK autoindex na każdej tabeli — lookup po record_id / id.", "")
    add_bullet(doc, "UNIQUE(crypto_id, snapshot_date) tworzy ukryty indeks złożony.", "")
    add_bullet(doc, "idx_snapshots_date — filtrowanie po dacie.", "")
    add_bullet(doc, "idx_current_collected_at — filtrowanie po czasie pobrania.", "")
    add_para(doc,
        "Schemat spełnia 3NF: metadane monety tylko w cryptocurrencies, "
        "tabele faktów zawierają wyłącznie crypto_id jako FK."
    )

    doc.add_page_break()

    # ══ 6. ARCHITEKTURA ═══════════════════════════════════════════════════════
    add_heading(doc, "6. Architektura systemu", 1)
    add_para(doc, "System oparty jest o trójwarstwową architekturę:")
    add_diagram(doc, rendered.get("architecture"),
                "Rys. 2. Architektura 3-warstwowa systemu", width_cm=14.0)
    add_diagram(doc, rendered.get("sequence"),
                "Rys. 3. Diagram sekwencji — pobieranie danych historycznych", width_cm=14.0)

    add_heading(doc, "6.1 Przepływ danych historycznych", 2)
    for s in [
        "Użytkownik inicjuje pobieranie (notebook Stage 5 lub Streamlit Data Collection).",
        "Dla każdej z 5 monet: GET /coins/{id}/market_chart?days=365.",
        "Konwersja timestamp_ms → YYYY-MM-DD (UTC).",
        "INSERT OR REPLACE do market_snapshots (executemany).",
        "Opóźnienie 10 s między monetami (rate limit API).",
    ]:
        add_numbered(doc, s)

    add_heading(doc, "6.2 Przepływ danych live", 2)
    for s in [
        "GET /coins/markets?ids=bitcoin,ethereum,solana,binancecoin,ripple.",
        "INSERT do market_current z collected_at = UTC now.",
    ]:
        add_numbered(doc, s)

    doc.add_page_break()

    # ══ 7. WARSTWA DANYCH ═════════════════════════════════════════════════════
    add_heading(doc, "7. Implementacja — warstwa danych", 1)

    add_heading(doc, "7.1 Inicjalizacja bazy", 2)
    add_para(doc,
        "Funkcja create_database() wywoływana przy starcie app.py i notebooka (Stage 3). "
        "Operacja idempotentna dzięki CREATE TABLE IF NOT EXISTS."
    )

    add_heading(doc, "7.2 Idempotentny zapis", 2)
    add_code(doc, """\
conn.executemany(
    "INSERT OR REPLACE INTO market_snapshots "
    "(crypto_id, snapshot_date, price_usd, market_cap, total_volume) "
    "VALUES (?,?,?,?,?)",
    rows,
)
conn.commit()""")

    add_heading(doc, "7.3 Odczyt do DataFrame", 2)
    add_code(doc,
        "@st.cache_data(ttl=60)\n"
        "def load_snapshots() -> pd.DataFrame:\n"
        "    df = pd.read_sql(\n"
        "        'SELECT s.snapshot_date, s.price_usd, s.market_cap,'\n"
        "        ' s.total_volume, c.name, c.symbol'\n"
        "        ' FROM market_snapshots s'\n"
        "        ' JOIN cryptocurrencies c ON s.crypto_id = c.id'\n"
        "        ' ORDER BY s.snapshot_date', conn)\n"
        "    return df"
    )

    doc.add_page_break()

    # ══ 8. WARSTWA PREZENTACJI ════════════════════════════════════════════════
    add_heading(doc, "8. Implementacja — warstwa prezentacji", 1)

    add_heading(doc, "8.1 Notebook Jupyter — 11 etapów", 2)
    add_table(doc,
        ["Etap", "Opis"],
        [
            ["1–2", "Środowisko uv, importy, stałe COINS / METRIC_MAP"],
            ["3–4", "DDL bazy, wstawienie listy monet"],
            ["5", "Pobieranie 365-dniowej historii + zapis do DB (~1826 wierszy)"],
            ["6–7", "Weryfikacja DB, wczytanie do pandas DataFrame"],
            ["8", "Szeregi czasowe — ipywidgets + Plotly"],
            ["9", "Analiza ilościowa — bar/box/violin"],
            ["10", "Dashboard rynkowy — KPI, heatmapa, treemap"],
            ["11", "Korelacja i zmienność — corr(), annualised volatility"],
        ],
        col_widths=[2.0, 15.5],
    )

    add_heading(doc, "8.2 Streamlit — 6 stron", 2)
    add_table(doc,
        ["Strona", "Opis"],
        [
            ["Overview", "Statystyki bazy, zakres dat, nawigacja"],
            ["Data Collection", "Pobieranie historyczne i live z CoinGecko"],
            ["Time Series", "Wykres liniowy — 6 filtrów sidebar"],
            ["Quantitative Analysis", "Bar / Box / Violin — 6 filtrów"],
            ["Market Dashboard", "KPI, tabela, grouped bar, heatmap, treemap"],
            ["Correlation & Volatility", "Macierz korelacji, zmienność, korelacja z BTC"],
        ],
        col_widths=[5.0, 12.5],
    )

    doc.add_page_break()

    # ══ 9. ANALIZY ════════════════════════════════════════════════════════════
    add_heading(doc, "9. Opis analiz i wizualizacji", 1)

    add_heading(doc, "9.1 Szeregi czasowe", 2)
    add_para(doc,
        "Wykres liniowy z opcjonalną średnią kroczącą (MA) i skalą logarytmiczną. "
        "Umożliwia porównanie monet o różnych cenach (BTC ~$90k vs XRP ~$2)."
    )

    add_heading(doc, "9.2 Analiza ilościowa", 2)
    add_para(doc,
        "Bar chart — porównanie agregacji. Box plot — mediana, kwartyle, outliers. "
        "Violin plot — dodatkowo gęstość rozkładu."
    )

    add_heading(doc, "9.3 Korelacja i zmienność", 2)
    add_code(doc, """\
returns = price_pivot.pct_change().dropna()
corr    = returns.corr()           # macierz Pearsona
vol     = returns.std() * (365**0.5) * 100   # zmienność roczna %""")

    doc.add_page_break()

    # ══ 10. DANE ════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Dane zebrane w projekcie", 1)
    add_table(doc,
        ["Tabela", "Wiersze", "Opis"],
        [
            ["cryptocurrencies", "5", "BTC, ETH, SOL, BNB, XRP"],
            ["market_snapshots", "1 826", "366 dni × 5 monet (2025-05-26 → 2026-05-26)"],
            ["market_current", "10", "2 pobrania live × 5 monet"],
        ],
        col_widths=[5.0, 2.5, 10.0],
    )

    # ══ 11. URUCHOMIENIE ════════════════════════════════════════════════════════
    add_heading(doc, "11. Instrukcja uruchomienia", 1)
    add_code(doc, """\
# Instalacja
uv sync

# Aplikacja webowa (zalecane)
uv run streamlit run app.py
# → http://localhost:8501

# Notebook
uv run jupyter lab crypto_market_analysis.ipynb""")
    add_para(doc,
        "Przy pierwszym uruchomieniu przejdź do Data Collection i pobierz dane historyczne."
    )

    doc.add_page_break()

    # ══ 12. PROBLEMY ══════════════════════════════════════════════════════════
    add_heading(doc, "12. Napotkane problemy i rozwiązania", 1)
    add_table(doc,
        ["Problem", "Przyczyna", "Rozwiązanie"],
        [
            ["ModuleNotFoundError", "Wymagany Python 3.14", "Zmiana na requires-python >= 3.13"],
            ["HTTP 429", "Rate limit CoinGecko", "REQUEST_DELAY = 10 s + retry"],
            ["Pusty dashboard", "market_current pusta", "Obliczenia z market_snapshots"],
            ["applymap deprecated", "pandas 2.1+", "Zamiana na .map()"],
        ],
        col_widths=[4.5, 5.5, 7.5],
    )

    doc.add_page_break()

    # ══ 13. WNIOSKI ═════════════════════════════════════════════════════════════
    add_heading(doc, "13. Wnioski", 1)

    add_heading(doc, "13.1 Osiągnięcia", 2)
    for title, desc in [
        ("System end-to-end", "API → SQLite → wizualizacje w Jupyter i Streamlit."),
        ("Poprawny schemat DB", "3NF, FK, UNIQUE, indeksy zoptymalizowane pod zapytania czasowe."),
        ("Idempotentny ETL", "INSERT OR REPLACE — bezpieczne ponowne pobieranie."),
        ("Reprodukowalność", "uv + uv.lock — identyczne środowisko na każdej maszynie."),
    ]:
        add_bullet(doc, desc, title + ":")

    add_heading(doc, "13.2 Możliwe rozszerzenia", 2)
    for e in [
        "Więcej monet — rozszerzenie listy COINS.",
        "Harmonogram — APScheduler / cron.",
        "PostgreSQL — dla większej skali i współbieżności.",
        "Eksport CSV/Excel — st.download_button w Streamlit.",
    ]:
        add_bullet(doc, e)

    add_heading(doc, "13.3 Uwagi końcowe", 2)
    add_para(doc,
        "SQLite w pełni wystarcza do lokalnych analiz — 1826 wierszy obsługiwanych "
        "poniżej 50 ms dzięki prawidłowo zaprojektowanym indeksom, w tym wykorzystaniu "
        "UNIQUE jako indeksu złożonego na (crypto_id, snapshot_date)."
    )

    doc.add_page_break()

    # ══ 14. LITERATURA ════════════════════════════════════════════════════════
    add_heading(doc, "14. Literatura i źródła", 1)
    for i, item in enumerate([
        "CoinGecko API Documentation — https://www.coingecko.com/en/api/documentation",
        "SQLite Documentation — https://www.sqlite.org/docs.html",
        "pandas Documentation — https://pandas.pydata.org/docs/",
        "Plotly Python — https://plotly.com/python/",
        "Streamlit Documentation — https://docs.streamlit.io/",
        "uv Documentation — https://docs.astral.sh/uv/",
        "E. F. Codd, A Relational Model of Data for Large Shared Data Banks, CACM, 1970.",
        "C. J. Date, An Introduction to Database Systems, Addison-Wesley, 8th ed., 2003.",
        "W. McKinney, Python for Data Analysis, O'Reilly, 3rd ed., 2022.",
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        r = p.add_run(f"[{i}]  ")
        r.bold = True
        r.font.size = Pt(10)
        r = p.add_run(item)
        r.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Sprawozdanie przygotowane w ramach przedmiotu Zaawansowane Bazy Danych\n"
        "Michał Dusza · Szymon Bugajski · Mateusz Basiura · Maj 2026"
    )
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUTPUT)
    print(f"\nZapisano: {OUTPUT}")

    for p in TMP_DIR.glob("*.png"):
        p.unlink()
    if TMP_DIR.exists() and not any(TMP_DIR.iterdir()):
        TMP_DIR.rmdir()


if __name__ == "__main__":
    build()
